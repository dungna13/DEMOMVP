"""
doc_generator.py — Dịch vụ tự động soạn thảo văn bản hành chính và chuyển đổi định dạng (Markdown, DOCX, PDF)
"""

import os
import subprocess
import logging
import json
import docx
from docx.shared import Pt, Inches
from typing import Dict, Any, Optional
from src.core.ai_service import _call_llm, is_ai_available
from prompts.system_prompt import DOC_EXTRACTION_PROMPT

logger = logging.getLogger(__name__)

DRAFTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "drafts")
TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "templates", "doc_templates")

# Đảm bảo thư mục nháp tồn tại
os.makedirs(DRAFTS_DIR, exist_ok=True)


def get_template_variables(template_type: str) -> list:
    """Trả về danh sách các biến cần thiết cho từng loại biểu mẫu."""
    if template_type == "don_khieu_nai":
        return [
            "dia_danh", "ngay", "thang", "nam", "tieu_de_khieu_nai",
            "co_quan_kinh_gui", "nguoi_khieu_nai", "cccd", "cccd_ngay", "cccd_noi_cap",
            "dia_chi", "so_dien_thoai", "doi_tuong_khieu_nai", "noi_dung_su_viec",
            "can_cu_phap_ly", "ly_do_trai_phap_luat", "yeu_cau_chi_tiet", "tài_liệu_kèm_theo"
        ]
    elif template_type == "cong_van_tra_loi":
        return [
            "co_quan_ban_hanh", "so_cong_van", "ky_hieu", "nguoi_nhan",
            "dia_danh", "ngay", "thang", "nam", "ngay_nhan_don",
            "noi_dung_phản_anh", "can_cu_phap_ly", "noi_dung_dieu_luat",
            "giai_quyet_chi_tiet", "nguoi_ky"
        ]
    elif template_type == "to_trinh_mau":
        return [
            "co_quan_trinh", "dia_danh", "ngay", "thang", "nam",
            "noi_dung_trinh", "cap_tren_kinh_gui", "ly_do_trinh",
            "can_cu_phap_ly", "su_can_thiet", "de_xuat_1",
            "phuong_an_trien_khai", "kinh_phi_thoi_gian", "nguoi_trinh_ky"
        ]
    return []


def extract_entities_for_doc(
    question: str,
    answer: str,
    template_type: str,
) -> Dict[str, str]:
    """Sử dụng LLM để trích xuất các biến điền vào mẫu văn bản từ ngữ cảnh hội thoại."""
    variables = get_template_variables(template_type)
    variables_json_desc = {v: f"giá trị cho {v}" for v in variables}
    
    if not is_ai_available():
        # Fallback default values
        from datetime import datetime
        now = datetime.now()
        defaults = {v: f"[Chưa điền {v}]" for v in variables}
        defaults.update({
            "dia_danh": "Hà Nội",
            "ngay": str(now.day),
            "thang": str(now.month),
            "nam": str(now.year),
        })
        return defaults

    system_prompt = DOC_EXTRACTION_PROMPT.format(
        variables_json=json.dumps(variables, ensure_ascii=False, indent=2)
    )

    user_content = f"Câu hỏi người dùng:\n{question}\n\nCâu trả lời pháp lý:\n{answer}"
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_content}
    ]
    
    extracted_text = _call_llm(messages, temperature=0.1, max_tokens=1000)
    
    extracted_data = {}
    if extracted_text:
        try:
            json_start = extracted_text.find("{")
            json_end = extracted_text.rfind("}") + 1
            if json_start >= 0 and json_end > json_start:
                extracted_data = json.loads(extracted_text[json_start:json_end])
        except Exception as e:
            logger.warning(f"[DocGen] Failed to parse extracted JSON: {e}")

    # Điền giá trị mặc định cho những trường thiếu
    from datetime import datetime
    now = datetime.now()
    for v in variables:
        if v not in extracted_data or not str(extracted_data[v]).strip():
            if v == "dia_danh":
                extracted_data[v] = "Hà Nội"
            elif v == "ngay":
                extracted_data[v] = str(now.day)
            elif v == "thang":
                extracted_data[v] = str(now.month)
            elif v == "nam":
                extracted_data[v] = str(now.year)
            else:
                extracted_data[v] = f"[{v.replace('_', ' ').title()}]"
                
    return extracted_data


def render_markdown_template(template_type: str, data: Dict[str, str]) -> str:
    """Điền dữ liệu vào mẫu file markdown tương ứng."""
    template_path = os.path.join(TEMPLATES_DIR, f"{template_type}.md")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Không tìm thấy file mẫu: {template_path}")
        
    with open(template_path, "r", encoding="utf-8") as f:
        content = f.read()
        
    # Thay thế các biến {{ bien }}
    for key, val in data.items():
        placeholder = "{{" + f" {key} " + "}}"
        placeholder_no_space = "{{" + key + "}}"
        content = content.replace(placeholder, str(val))
        content = content.replace(placeholder_no_space, str(val))
        
    return content


def markdown_to_docx(md_content: str, docx_path: str):
    """Chuyển đổi văn bản Markdown sang tệp DOCX chuẩn Nghị định 30/2020/NĐ-CP."""
    doc = docx.Document()
    
    # Thiết lập căn lề chuẩn hành chính Việt Nam (Top: 20mm, Bottom: 20mm, Left: 30mm, Right: 15mm)
    for section in doc.sections:
        section.top_margin = Inches(0.79)     # ~20mm
        section.bottom_margin = Inches(0.79)  # ~20mm
        section.left_margin = Inches(1.18)    # ~30mm
        section.right_margin = Inches(0.59)   # ~15mm

    # Thiết lập font chữ mặc định là Times New Roman, cỡ chữ 13pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Xử lý từng dòng văn bản
    lines = md_content.split('\n')
    for line in lines:
        stripped = line.strip()
        
        # Tiêu đề
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(14)
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.size = Pt(13)
        # Gạch đầu dòng / Danh sách
        elif stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        # Đường gạch ngang dòng ký tự
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("—" * 30)
            run.font.size = Pt(10)
        else:
            p = doc.add_paragraph(stripped)
            # Cấu hình dòng căn lề cho dòng "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
            if "CỘNG HÒA XÃ HỘI" in stripped or "Độc lập - Tự do" in stripped or "ĐƠN KHIẾU NẠI" in stripped or "TỜ TRÌNH" in stripped:
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                # Tăng in đậm cho tiêu đề đơn
                if "ĐƠN KHIẾU NẠI" in stripped or "TỜ TRÌNH" in stripped:
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(14)
            # Căn lề phải cho ngày tháng địa danh
            elif ", ngày" in stripped and "tháng" in stripped and "năm" in stripped:
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                
    doc.save(docx_path)


def convert_docx_to_pdf(docx_path: str, output_dir: str) -> Optional[str]:
    """Chuyển đổi tệp DOCX sang PDF bằng LibreOffice Headless."""
    # Đường dẫn cài đặt LibreOffice mặc định trên Windows
    soffice_paths = [
        "C:\\Program Files\\LibreOffice\\program\\soffice.com",
        "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        "soffice"
    ]
    
    soffice_bin = None
    for path in soffice_paths:
        try:
            # Kiểm tra xem có thể chạy được không
            subprocess.run([path, "--version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
            soffice_bin = path
            break
        except Exception:
            continue
            
    if not soffice_bin:
        logger.warning("[DocGen] LibreOffice is not installed or not in PATH. PDF conversion skipped.")
        return None

    try:
        cmd = [
            soffice_bin,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", output_dir,
            docx_path
        ]
        logger.info(f"[DocGen] Converting DOCX to PDF: {' '.join(cmd)}")
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True, timeout=30)
        
        # Tên tệp PDF sinh ra
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(output_dir, f"{base_name}.pdf")
        
        if os.path.exists(pdf_path):
            logger.info(f"[DocGen] PDF successfully generated at {pdf_path}")
            return pdf_path
            
    except Exception as e:
        logger.error(f"[DocGen] PDF conversion failed: {e}")
        
    return None


def generate_document_draft(
    question: str,
    answer: str,
    template_type: str,
    draft_id: str
) -> Dict[str, str]:
    """Sinh toàn bộ các định dạng văn bản (MD, DOCX, PDF) từ biểu mẫu hành chính."""
    # Bước 1: Trích xuất các biến
    data = extract_entities_for_doc(question, answer, template_type)
    
    # Bước 2: Sinh nội dung Markdown
    md_content = render_markdown_template(template_type, data)
    md_path = os.path.join(DRAFTS_DIR, f"{draft_id}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
        
    # Bước 3: Sinh file DOCX
    docx_path = os.path.join(DRAFTS_DIR, f"{draft_id}.docx")
    markdown_to_docx(md_content, docx_path)
    
    # Bước 4: Chuyển đổi sang PDF
    pdf_path = convert_docx_to_pdf(docx_path, DRAFTS_DIR)
    
    return {
        "draft_id": draft_id,
        "md_path": md_path,
        "docx_path": docx_path,
        "pdf_path": pdf_path if pdf_path else "",
        "data": data,
        "content_preview": md_content
    }
